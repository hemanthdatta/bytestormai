import os
import io
import re
import traceback
import google.generativeai as genai
from PIL import Image
import pandas as pd
from collections import Counter

# For DOCX processing
from docx import Document
from docx.document import Document as DocumentClass
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.section import CT_SectPr
from docx.oxml import OxmlElement

# For PPTX processing
from pptx import Presentation
from pptx.enum.shapes import MSO_SHAPE_TYPE

from pdf2md import summarize_markdown_groups

def create_safe_filename(text):
    """Create a safe filename from text by removing invalid characters."""
    return re.sub(r'[\\/*?:"<>|]', "", text.replace(' ', '_'))[:50]

def get_font_size_stats(doc):
    """
    Scan the entire DOCX to collect and count all font sizes.
    """
    font_sizes = []
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            font_size = run.font.size
            if font_size:
                # Convert to points if it's in twips (1/20th of a point)
                if font_size.pt:
                    font_sizes.append(round(font_size.pt))
                elif hasattr(font_size, 'val') and font_size.val:
                    # Convert from twips to points
                    font_sizes.append(round(font_size.val / 20))
    return Counter(font_sizes)

def get_heading_thresholds(font_counter):
    """
    Determine thresholds for section (##) and subsection (###) headings.
    Most frequent size = body text; largest above that = section; next = subsection.
    """
    if not font_counter:
        return 0, 0
    sizes = sorted(font_counter.keys(), reverse=True)
    body_size = font_counter.most_common(1)[0][0]
    section_size = sizes[0] if sizes[0] > body_size else body_size + 1
    subsection_size = (sizes[1] if len(sizes) > 1 and sizes[1] > body_size else body_size + 0.5)
    return section_size, subsection_size

def generate_md(docx_path, api_key, output_file=None, max_workers=4):
    """
    Extract content from a DOCX file into a markdown file with images and tables.
    
    Args:
        docx_path (str): Path to the DOCX file
        api_key (str): Google API key for Gemini
        output_file (str): Optional output file name (without .md extension)
        max_workers (int): Maximum number of parallel workers (not used for DOCX, kept for API compatibility)
    """
    try:
        # Configure the Gemini API
        genai.configure(api_key=api_key)
        
        # Check if file exists
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")
            
        # Setup output directory and file
        docx_name = os.path.splitext(os.path.basename(docx_path))[0]
        
        if output_file is None:
            output_file = docx_name
            output_dir = os.path.join(os.path.dirname(docx_path), f"{docx_name}_markdown")
        else:
            output_dir = os.path.dirname(os.path.abspath(output_file))
            if not output_dir:  # If output_file has no directory component
                output_dir = os.path.join(os.path.dirname(docx_path), f"{output_file}_markdown")
        
        # Create output directories
        images_dir = os.path.join(output_dir, "images")
        table_images_dir = os.path.join(output_dir, "table_images")
        os.makedirs(images_dir, exist_ok=True)
        os.makedirs(table_images_dir, exist_ok=True)
        
        # Initialize markdown content
        markdown_content = f"# {docx_name}\n\n"
        
        # Initialize Gemini model - with error handling
        try:
            model = genai.GenerativeModel('gemini-2.0-flash')
        except Exception as e:
            print(f"Error initializing Gemini model: {str(e)}")
            print("Trying with model name 'gemini-pro-vision'...")
            try:
                model = genai.GenerativeModel('gemini-pro-vision')
            except Exception as e2:
                print(f"Error with alternate model: {str(e2)}")
                raise Exception("Failed to initialize Gemini model. Check your API key and network connection.")
        
        # Open the DOCX with error handling
        try:
            doc = Document(docx_path)
        except Exception as e:
            raise Exception(f"Failed to open DOCX: {str(e)}")
        
        # Get font statistics and heading thresholds
        font_counter = get_font_size_stats(doc)
        section_size, subsection_size = get_heading_thresholds(font_counter)
        print(f"DOCX analysis: section_size={section_size}, subsection_size={subsection_size}")
        
        # Process document
        current_page = 1
        all_elements = list(iter_block_items(doc))  # Pre-load all elements
        page_breaks = find_page_breaks(doc)
        
        # Debug information
        print(f"Found {len(page_breaks)} page breaks at indices: {page_breaks}")
        
        # Dictionary to track images that have been processed
        processed_image_ids = set()
        image_count = 0
        
        # Iterate through all elements in the document
        for element_count, element in enumerate(all_elements):
            # Check if we've hit a page break
            if element_count in page_breaks:
                current_page += 1
                markdown_content += f"\n\n---\n\n## Page {current_page}\n\n"
            
            if isinstance(element, Paragraph):
                # Process paragraph
                if element.text.strip():
                    # Format paragraph text based on styles and font sizes
                    if element.style.name.startswith('Heading'):
                        level = int(element.style.name.replace('Heading', '')) if element.style.name != 'Heading' else 1
                        level = min(max(level, 1), 6)  # Ensure level is between 1-6
                        markdown_content += f"{'#' * level} {element.text.strip()}\n\n"
                    else:
                        # Check if paragraph might be a heading based on font size
                        max_size = 0
                        for run in element.runs:
                            if run.font.size:
                                # Get font size in points
                                if run.font.size.pt:
                                    size = run.font.size.pt
                                elif hasattr(run.font.size, 'val') and run.font.size.val:
                                    size = run.font.size.val / 20  # Convert from twips
                                else:
                                    size = 0
                                max_size = max(max_size, size)
                        
                        if max_size >= section_size:
                            markdown_content += f"## {element.text.strip()}\n\n"
                        elif max_size >= subsection_size:
                            markdown_content += f"### {element.text.strip()}\n\n"
                        else:
                            markdown_content += f"{element.text.strip()}\n\n"
                
                # Check for inline images in this paragraph
                for run in element.runs:
                    if run._r.drawing_lst:
                        for drawing in run._r.drawing_lst:
                            # Try to get the embedded image
                            try:
                                # Access the blip relationship ID without using xpath
                                blip = None
                                # Navigate through the XML structure to find blip elements
                                for inline_or_anchor in drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline') or drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor'):
                                    for graphic in inline_or_anchor.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}graphic'):
                                        for graphicData in graphic.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}graphicData'):
                                            for pic in graphicData.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic'):
                                                for blipFill in pic.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}blipFill'):
                                                    blips = blipFill.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                                                    if blips:
                                                        blip = blips[0]
                                                        break
                                
                                if blip is not None:
                                    # Get the relationship ID
                                    image_rel_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                    if image_rel_id and image_rel_id not in processed_image_ids:
                                        processed_image_ids.add(image_rel_id)
                                        image_count += 1
                                        
                                        # Get image data from relationship
                                        image_part = doc.part.related_parts[image_rel_id]
                                        image_data = image_part.blob
                                        
                                        # Get image extension
                                        ext = image_part.partname.ext
                                        if not ext:
                                            ext = '.png'  # Default to PNG if no extension
                                        
                                        # Save the image
                                        img_filename = f"image_{image_count}{ext}"
                                        img_path = os.path.join(images_dir, img_filename)
                                        
                                        with open(img_path, "wb") as img_file:
                                            img_file.write(image_data)
                                        
                                        # Add image to markdown
                                        markdown_content += f"### Image {image_count}\n\n"
                                        markdown_content += f"![Image {image_count}](images/{img_filename})\n\n"
                                        
                                        # Use Gemini for OCR on the image
                                        try:
                                            image = Image.open(io.BytesIO(image_data))
                                            response = model.generate_content(["Extract all text from this image", image])
                                            ocr_text = response.text
                                            if ocr_text and ocr_text.strip():
                                                markdown_content += f"**OCR Text from Image {image_count}:**\n\n{ocr_text}\n\n"
                                        except Exception as e:
                                            markdown_content += f"**OCR Failed for Image {image_count}:** {str(e)}\n\n"
                            except Exception as e:
                                print(f"Error processing inline image: {str(e)}")
            
            elif isinstance(element, Table):
                # Process table
                markdown_content += "### Table\n\n"
                
                # Convert table to markdown
                table_md = ""
                
                # Get all cells and determine number of rows and columns
                rows = len(element.rows)
                cols = len(element.columns)
                
                # Create table header
                table_md += "| " + " | ".join(["Column " + str(i+1) for i in range(cols)]) + " |\n"
                table_md += "| " + " | ".join(["---" for _ in range(cols)]) + " |\n"
                
                # Fill in table data
                for i, row in enumerate(element.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                        # Replace any pipe symbols to avoid breaking markdown tables
                        cell_text = cell_text.replace("|", "\\|")
                        if not cell_text:
                            cell_text = "-"
                        row_data.append(cell_text)
                    table_md += "| " + " | ".join(row_data) + " |\n"
                
                markdown_content += table_md + "\n\n"
                
                # Try to extract table as pandas DataFrame for more complex tables
                try:
                    # Create a more structured table representation
                    data = []
                    for row in element.rows:
                        data_row = []
                        for cell in row.cells:
                            cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                            data_row.append(cell_text if cell_text else "-")
                        data.append(data_row)
                    
                    # Create DataFrame (first row as header)
                    if len(data) > 1:
                        df = pd.DataFrame(data[1:], columns=data[0])
                        # Add DataFrame rendering as CSV
                        markdown_content += "#### Table as CSV\n\n```\n"
                        markdown_content += df.to_csv(index=False)
                        markdown_content += "```\n\n"
                except Exception as e:
                    print(f"Failed to convert table to DataFrame: {str(e)}")
        
        # Process any remaining images not found inline (like floating images)
        for rel in doc.part.rels.values():
            if "image" in rel.reltype and rel.rId not in processed_image_ids:
                try:
                    image_count += 1
                    # Get image data
                    image_data = rel.target_part.blob
                    
                    # Save the image
                    ext = rel.target_ref.split(".")[-1] if "." in rel.target_ref else "png"
                    img_filename = f"image_{image_count}.{ext}"
                    img_path = os.path.join(images_dir, img_filename)
                    
                    with open(img_path, "wb") as img_file:
                        img_file.write(image_data)
                    
                    # Add image to markdown with a note that it's a floating image
                    markdown_content += f"### Floating Image {image_count}\n\n"
                    markdown_content += f"![Floating Image {image_count}](images/{img_filename})\n\n"
                    
                    # Use Gemini for OCR on the image
                    try:
                        image = Image.open(io.BytesIO(image_data))
                        response = model.generate_content(["Extract all text from this image", image])
                        ocr_text = response.text
                        if ocr_text and ocr_text.strip():
                            markdown_content += f"**OCR Text from Image {image_count}:**\n\n{ocr_text}\n\n"
                    except Exception as e:
                        markdown_content += f"**OCR Failed for Image {image_count}:** {str(e)}\n\n"
                    
                except Exception as e:
                    markdown_content += f"**Error processing Image {image_count}:** {str(e)}\n\n"
        
        # Save the final markdown file
        markdown_file = f"{output_file}.md"
        with open(markdown_file, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        
        print(f"Markdown file saved to: {markdown_file}")
        
    except Exception as e:
        print(f"Error occurred: {str(e)}")
        print(traceback.format_exc())

def find_page_breaks(doc):
    """
    Find approximate page breaks in the document.
    
    This function attempts to find page breaks based on explicit page break
    elements, section breaks, and paragraph properties.
    
    Args:
        doc: The Document object
        
    Returns:
        A list of indices where page breaks occur
    """
    page_breaks = set()  # Using a set to avoid duplicates
    
    # First pass: Check all paragraphs and their runs for explicit page breaks
    all_blocks = list(iter_block_items(doc))
    
    for idx, block in enumerate(all_blocks):
        if isinstance(block, Paragraph):
            # Check for explicit page breaks in runs
            for run in block.runs:
                if run.element.tag.endswith('br') and run.element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                    page_breaks.add(idx)
            
            # Check for section breaks
            if block._p.getparent() is not None:
                last_children = block._p.getparent().getchildren()
                if last_children and last_children[-1].tag.endswith('sectPr'):
                    sect_pr = last_children[-1]
                    # Check if there's a type attribute indicating a page break
                    type_elements = sect_pr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                    if type_elements:
                        type_val = type_elements[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                        if type_val in ['nextPage', 'evenPage', 'oddPage']:
                            page_breaks.add(idx)
            
            # Check for page break before in paragraph properties
            if hasattr(block._p, 'pPr') and block._p.pPr is not None and hasattr(block._p.pPr, 'pageBreakBefore') and block._p.pPr.pageBreakBefore is not None:
                page_breaks.add(idx)
                
            # Check for w:br with type='page' in paragraph
            try:
                for br in block._p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'):
                    if br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                        page_breaks.add(idx)
                        break
            except AttributeError:
                # Skip if findall is not available
                pass
    
    # Check for explicit w:lastRenderedPageBreak elements (Word's pagination info)
    for idx, block in enumerate(all_blocks):
        if isinstance(block, Paragraph):
            try:
                if block._p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lastRenderedPageBreak'):
                    page_breaks.add(idx)
            except AttributeError:
                # Skip if findall is not available
                pass
    
    return sorted(list(page_breaks))

# Helper function to iterate through blocks in a Word document
def iter_block_items(parent):
    """
    Generate a sequence of block-level items in a Word document.
    This is more reliable than using parent.paragraphs and parent.tables,
    as it maintains the original order of the elements.
    """
    if isinstance(parent, DocumentClass):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)